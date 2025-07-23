from crewai import Agent, Task
import os
from tools.document_tools import load_document, save_document

try:
    from docx import Document
except ImportError:
    Document = None  # python-docx not installed

class AssemblyAgent:
    """
    crewAI-powered agent for assembling the final BRSR report from generated content sections.
    Uses deterministic code for document assembly (Word, extendable to PDF, etc.).
    """
    def __init__(self, gemini_api_key: str = None):
        self.gemini_api_key = gemini_api_key or os.getenv("GEMINI_API_KEY")
        self.agent = Agent(
            role="BRSR Report Assembler",
            goal="Assemble the final BRSR report from generated content sections and template."
        )

    def assemble_report_tool(self, content_sections: dict, template_path: str, output_path: str) -> str:
        if template_path.lower().endswith('.docx'):
            if Document is None:
                raise ImportError("python-docx is required for Word document assembly. Please install it.")
            doc = Document(template_path)
            for section_title, section_content in content_sections.items():
                found = False
                for para in doc.paragraphs:
                    if section_title.strip().lower() in para.text.strip().lower():
                        para.text = f"{section_title}\n{section_content}"
                        found = True
                        break
                if not found:
                    doc.add_paragraph(f"{section_title}\n{section_content}")
            doc.save(output_path)
            return f"Report assembled and saved to {output_path}"
        else:
            raise NotImplementedError("Only .docx templates are currently supported. Extend for PDF/other formats.")

    def get_task(self, content_sections: dict, template_path: str, output_path: str) -> Task:
        """
        Returns a crewAI Task for assembling the final report.
        """
        return Task(
            description="Assemble the final BRSR report from generated content sections and template.",
            agent=self.agent,
            expected_output="A message indicating the report was assembled and saved.",
            func=lambda: self.assemble_report_tool(content_sections, template_path, output_path)
        ) 